import whisper
import time
import re
import os
from docx import Document
from fuzzywuzzy import fuzz
from typing import List, Tuple, Dict
import logging
import difflib

logger = logging.getLogger(__name__)

class AudioAnalyzer:
    def __init__(self, model_size='tiny'):
        """
        Initialize the audio analyzer with Whisper model
        model_size options: 'tiny', 'base', 'small', 'medium', 'large'
        For large files, 'tiny' is recommended for speed
        """
        self.model = whisper.load_model(model_size)
        self.model_size = model_size
    
    def get_audio_duration(self, audio_path: str) -> float:
        """Get audio duration in seconds"""
        try:
            import librosa
            duration = librosa.get_duration(path=audio_path)
            return duration
        except Exception as e:
            logger.warning(f"Could not get audio duration: {e}")
            return 0
    
    def estimate_processing_time(self, audio_path: str) -> float:
        """Estimate processing time based on audio duration and model size"""
        duration = self.get_audio_duration(audio_path)
        
        # Rough estimates based on model size (seconds per minute of audio)
        model_speeds = {
            'tiny': 2.0,    # 2 seconds per minute
            'base': 4.0,    # 4 seconds per minute
            'small': 8.0,   # 8 seconds per minute
            'medium': 16.0, # 16 seconds per minute
            'large': 32.0   # 32 seconds per minute
        }
        
        speed = model_speeds.get(self.model_size, 4.0)
        estimated_time = (duration / 60) * speed
        
        return estimated_time
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = Document(docx_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            return ' '.join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def transcribe_audio(self, audio_path: str) -> str:
        """Transcribe audio file using Whisper with optimized settings"""
        try:
            start_time = time.time()
            
            # Check file size and warn if too large
            file_size = os.path.getsize(audio_path) / (1024 * 1024)  # MB
            if file_size > 100:
                logger.warning(f"Large file detected: {file_size:.1f}MB. This may take a long time.")
                print(f"[WARNING] Large file: {file_size:.1f}MB. Processing may take a long time or fail. For best results, split into smaller chunks.")
            
            # Try multiple approaches to avoid ffmpeg issues
            try:
                # First try: Direct Whisper (might work if ffmpeg is available)
                logger.info("Attempting direct Whisper transcription")
                result = self.model.transcribe(
                    audio_path,
                    fp16=False,
                    language=None,
                    task="transcribe",
                    verbose=False,
                    condition_on_previous_text=False,
                    temperature=0.0
                )
                logger.info("Direct Whisper transcription successful")
                
            except Exception as direct_error:
                logger.warning(f"Direct Whisper failed: {direct_error}")
                
                # Second try: Use librosa if available
                try:
                    import librosa
                    import soundfile as sf
                    
                    logger.info("Using librosa for audio loading")
                    
                    # Load audio with librosa and resample to 16kHz
                    audio, sr = librosa.load(audio_path, sr=16000)
                    
                    # Save as temporary WAV file for Whisper
                    temp_path = audio_path + "_temp.wav"
                    sf.write(temp_path, audio, sr)
                    
                    logger.info(f"Created temporary file: {temp_path}")
                    
                    # Use the temporary file for transcription
                    result = self.model.transcribe(
                        temp_path,
                        fp16=False,
                        language=None,
                        task="transcribe",
                        verbose=False,
                        condition_on_previous_text=False,
                        temperature=0.0
                    )
                    
                    # Clean up temporary file
                    try:
                        os.remove(temp_path)
                        logger.info("Cleaned up temporary file")
                    except Exception as cleanup_error:
                        logger.warning(f"Could not clean up temp file: {cleanup_error}")
                        
                except Exception as librosa_error:
                    logger.warning(f"librosa failed: {librosa_error}")
                    
                    # Third try: Manual audio loading with wave/numpy
                    try:
                        import wave
                        import numpy as np
                        
                        logger.info("Using manual audio loading with wave/numpy")
                        
                        with wave.open(audio_path, 'rb') as wav_file:
                            frames = wav_file.getnframes()
                            sample_rate = wav_file.getframerate()
                            audio_data = wav_file.readframes(frames)
                            audio_array = np.frombuffer(audio_data, dtype=np.int16)
                            audio_float = audio_array.astype(np.float32) / 32768.0
                        
                        # Save as temporary file
                        temp_path = audio_path + "_numpy_temp.wav"
                        with wave.open(temp_path, 'wb') as temp_wav:
                            temp_wav.setnchannels(1)
                            temp_wav.setsampwidth(2)
                            temp_wav.setframerate(sample_rate)
                            temp_wav.writeframes(audio_float.tobytes())
                        
                        result = self.model.transcribe(
                            temp_path,
                            fp16=False,
                            language=None,
                            task="transcribe",
                            verbose=False,
                            condition_on_previous_text=False,
                            temperature=0.0
                        )
                        
                        # Clean up
                        try:
                            os.remove(temp_path)
                        except:
                            pass
                            
                    except Exception as numpy_error:
                        logger.error(f"All audio loading methods failed: {numpy_error}")
                        raise Exception("Could not load audio file with any available method")
            
            processing_time = time.time() - start_time
            
            logger.info(f"Transcription completed in {processing_time:.2f} seconds using {self.model_size} model")
            return result["text"].strip(), processing_time
        except Exception as e:
            logger.error(f"Error transcribing audio: {e}")
            raise
    
    def preprocess_text(self, text: str) -> List[str]:
        """Clean and tokenize text for comparison"""
        # Remove extra whitespace and convert to lowercase
        text = re.sub(r'\s+', ' ', text.lower().strip())
        # Remove punctuation except apostrophes
        text = re.sub(r'[^\w\s\']', '', text)
        # Split into words
        words = text.split()
        return words
    
    def calculate_similarity(self, word1: str, word2: str) -> float:
        """Calculate similarity between two words using fuzzy matching"""
        return fuzz.ratio(word1.lower(), word2.lower())
    
    def align_texts(self, script_words: List[str], audio_words: List[str]) -> List[Dict]:
        """Align script and audio words using SequenceMatcher for optimal alignment"""
        comparisons = []
        matcher = difflib.SequenceMatcher(None, script_words, audio_words, autojunk=False)
        opcodes = matcher.get_opcodes()
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'equal':
                for idx in range(i2 - i1):
                    comparisons.append({
                        'script_word': script_words[i1 + idx],
                        'audio_word': audio_words[j1 + idx],
                        'word_index': i1 + idx,
                        'is_correct': True,
                        'similarity_score': 100,
                        'error_type': 'correct'
                    })
            elif tag == 'replace':
                for idx in range(max(i2 - i1, j2 - j1)):
                    s_word = script_words[i1 + idx] if i1 + idx < i2 else ''
                    a_word = audio_words[j1 + idx] if j1 + idx < j2 else ''
                    similarity = self.calculate_similarity(s_word, a_word) if s_word and a_word else 0
                    comparisons.append({
                        'script_word': s_word,
                        'audio_word': a_word,
                        'word_index': i1 + idx,
                        'is_correct': similarity >= 80,
                        'similarity_score': similarity,
                        'error_type': 'wrong' if s_word and a_word else ('missing' if s_word else 'extra')
                    })
            elif tag == 'delete':
                for idx in range(i1, i2):
                    comparisons.append({
                        'script_word': script_words[idx],
                        'audio_word': '',
                        'word_index': idx,
                        'is_correct': False,
                        'similarity_score': 0,
                        'error_type': 'missing'
                    })
            elif tag == 'insert':
                for idx in range(j1, j2):
                    comparisons.append({
                        'script_word': '',
                        'audio_word': audio_words[idx],
                        'word_index': i1,
                        'is_correct': False,
                        'similarity_score': 0,
                        'error_type': 'extra'
                    })
        return comparisons
    
    def analyze_audio_accuracy(self, script_path: str, audio_path: str) -> Dict:
        """Main analysis function with performance optimizations"""
        try:
            # Get file info for estimation
            file_size = os.path.getsize(audio_path) / (1024 * 1024)  # MB
            duration = self.get_audio_duration(audio_path)
            estimated_time = self.estimate_processing_time(audio_path)
            
            logger.info(f"Starting analysis: {file_size:.1f}MB, {duration:.1f}s, estimated time: {estimated_time:.1f}s")
            
            # Extract script text
            script_text = self.extract_text_from_docx(script_path)
            script_words = self.preprocess_text(script_text)
            
            # Transcribe audio
            transcribed_text, processing_time = self.transcribe_audio(audio_path)
            audio_words = self.preprocess_text(transcribed_text)
            
            # Align and compare texts
            comparisons = self.align_texts(script_words, audio_words)
            
            # Calculate statistics
            total_words = len(script_words)
            correct_words = sum(1 for comp in comparisons if comp['is_correct'])
            missing_words = sum(1 for comp in comparisons if comp['error_type'] == 'missing')
            wrong_words = sum(1 for comp in comparisons if comp['error_type'] == 'wrong')
            
            accuracy_score = (correct_words / total_words * 100) if total_words > 0 else 0
            
            logger.info(f"Analysis completed: {accuracy_score:.1f}% accuracy, {processing_time:.1f}s actual time")
            
            return {
                'script_text': script_text,
                'transcribed_text': transcribed_text,
                'processing_time': processing_time,
                'file_info': {
                    'size_mb': file_size,
                    'duration_seconds': duration,
                    'estimated_time': estimated_time
                },
                'comparisons': comparisons,
                'statistics': {
                    'total_words': total_words,
                    'correct_words': correct_words,
                    'missing_words': missing_words,
                    'wrong_words': wrong_words,
                    'accuracy_score': accuracy_score
                }
            }
            
        except Exception as e:
            logger.error(f"Error in audio analysis: {e}")
            raise
    
    def extract_paragraphs_from_docx(self, docx_path: str) -> list:
        """Extract paragraphs from a DOCX file as a list of strings."""
        try:
            doc = Document(docx_path)
            return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception as e:
            logger.error(f"Error extracting paragraphs from DOCX: {e}")
            return []

    def split_text_into_paragraphs(self, text: str) -> list:
        """Split text into paragraphs by period or newlines."""
        import re
        # Split on period, question mark, exclamation, or newlines
        paras = re.split(r'[.!?]\s+|\n+', text)
        return [p.strip() for p in paras if p.strip()]

    def highlight_differences(self, script_para, audio_para):
        import difflib
        script_words = script_para.split()
        audio_words = audio_para.split()
        matcher = difflib.SequenceMatcher(None, script_words, audio_words, autojunk=False)
        script_result = []
        audio_result = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                for idx in range(i2 - i1):
                    script_result.append(script_words[i1 + idx])
                    audio_result.append(audio_words[j1 + idx])
            elif tag == 'replace':
                for idx in range(i2 - i1):
                    if i1 + idx < len(script_words):
                        script_result.append(f'<span class="wrong-word">{script_words[i1 + idx]}</span>')
                for idx in range(j2 - j1):
                    if j1 + idx < len(audio_words):
                        audio_result.append(f'<span class="wrong-word">{audio_words[j1 + idx]}</span>')
            elif tag == 'delete':
                for idx in range(i1, i2):
                    script_result.append(f'<span class="missing-word">{script_words[idx]}</span>')
            elif tag == 'insert':
                for idx in range(j1, j2):
                    audio_result.append(f'<span class="extra-word">{audio_words[idx]}</span>')
        return ' '.join(script_result), ' '.join(audio_result)

    def compare_paragraphs(self, script_text: str, audio_text: str, script_path: str = None, force_single_paragraph: bool = False) -> list:
        import difflib
        if force_single_paragraph:
            script_paragraphs = [script_text.strip()] if script_text.strip() else []
            audio_paragraphs = [audio_text.strip()] if audio_text.strip() else []
        else:
            if script_path:
                script_paragraphs = self.extract_paragraphs_from_docx(script_path)
            else:
                script_paragraphs = self.split_text_into_paragraphs(script_text)
            audio_paragraphs = self.split_text_into_paragraphs(audio_text)
        logger.error(f"[DEBUG] Script paragraphs ({len(script_paragraphs)}): {script_paragraphs}")
        logger.error(f"[DEBUG] Audio paragraphs ({len(audio_paragraphs)}): {audio_paragraphs}")
        print(f"[DEBUG] Script paragraphs ({len(script_paragraphs)}): {script_paragraphs}")
        print(f"[DEBUG] Audio paragraphs ({len(audio_paragraphs)}): {audio_paragraphs}")
        matcher = difflib.SequenceMatcher(None, script_paragraphs, audio_paragraphs, autojunk=False)
        opcodes = matcher.get_opcodes()
        results = []
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'equal':
                for idx in range(i2 - i1):
                    s_para = script_paragraphs[i1 + idx]
                    a_para = audio_paragraphs[j1 + idx]
                    highlighted_script, highlighted_audio = self.highlight_differences(s_para, a_para)
                    results.append({
                        'script_paragraph': highlighted_script,
                        'audio_paragraph': highlighted_audio,
                        'similarity': 100,
                        'status': 'Correct',
                        'index': i1 + idx
                    })
            elif tag == 'replace' or tag == 'delete' or tag == 'insert':
                for idx in range(max(i2 - i1, j2 - j1)):
                    s_para = script_paragraphs[i1 + idx] if i1 + idx < i2 else ''
                    a_para = audio_paragraphs[j1 + idx] if j1 + idx < j2 else ''
                    similarity = self.calculate_similarity(s_para, a_para) if s_para and a_para else 0
                    highlighted_script, highlighted_audio = self.highlight_differences(s_para, a_para)
                    results.append({
                        'script_paragraph': highlighted_script,
                        'audio_paragraph': highlighted_audio,
                        'similarity': similarity,
                        'status': 'Wrong' if s_para and a_para else ('Missing' if s_para else 'Extra'),
                        'index': i1 + idx
                    })
        logger.error(f"[DEBUG] Paragraph comparison results: {results}")
        print(f"[DEBUG] Paragraph comparison results: {results}")
        return results 