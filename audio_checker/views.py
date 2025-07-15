from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
import os
import threading
import time
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

from .models import AudioAnalysis, WordComparison, AnalysisResult, BatchUpload
from .forms import AudioAnalysisForm, BatchUploadForm
from .services import AudioAnalyzer

logger = logging.getLogger(__name__)

def home(request):
    """Home page with upload form"""
    if request.method == 'POST':
        form = AudioAnalysisForm(request.POST, request.FILES)
        if form.is_valid():
            analysis = form.save(commit=False)
            if request.user.is_authenticated:
                analysis.user = request.user
            analysis.save()
            
            # Start analysis in background with timeout
            thread = threading.Thread(target=run_analysis_with_timeout, args=(analysis.id,))
            thread.daemon = True
            thread.start()
            
            messages.success(request, 'Analysis started! You will be notified when it completes.')
            return redirect('audio_checker:analysis_detail', analysis_id=analysis.id)
    else:
        form = AudioAnalysisForm()
    
    return render(request, 'audio_checker/home.html', {'form': form})

def run_analysis_with_timeout(analysis_id):
    """Background task to run audio analysis with timeout and better error handling"""
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f'[BATCH DEBUG] Starting analysis for ID {analysis_id}')
    try:
        analysis = AudioAnalysis.objects.get(id=analysis_id)
        logger.info(f'[BATCH DEBUG] Loaded analysis: {analysis}')
        
        # Choose model size based on file size
        file_size = analysis.audio_file.size / (1024 * 1024)  # MB
        logger.info(f'[BATCH DEBUG] Audio file size: {file_size:.2f} MB')
        if file_size > 50:
            model_size = 'tiny'  # Use tiny model for large files
        elif file_size > 20:
            model_size = 'base'  # Use base model for medium files
        else:
            model_size = 'small'  # Use small model for small files
        logger.info(f'[BATCH DEBUG] Using model size: {model_size}')
        
        # Set a timeout for the entire analysis (30 minutes max)
        start_time = time.time()
        timeout = 1800  # 30 minutes
        
        try:
            analyzer = AudioAnalyzer(model_size=model_size)
            
            # Get file paths
            script_path = analysis.script_file.path
            audio_path = analysis.audio_file.path
            logger.info(f'[BATCH DEBUG] Script path: {script_path}, Audio path: {audio_path}')
            logger.info(f'[BATCH DEBUG] Script exists: {os.path.exists(script_path)}, Audio exists: {os.path.exists(audio_path)}')
            
            # Check if files exist
            if not os.path.exists(script_path):
                logger.error(f'[BATCH DEBUG] Script file not found: {script_path}')
                raise FileNotFoundError(f"Script file not found: {script_path}")
            if not os.path.exists(audio_path):
                logger.error(f'[BATCH DEBUG] Audio file not found: {audio_path}')
                raise FileNotFoundError(f"Audio file not found: {audio_path}")
            
            # Run analysis
            result = analyzer.analyze_audio_accuracy(script_path, audio_path)
            logger.info(f'[BATCH DEBUG] Analysis result: {result}')
            
            # Check timeout
            if time.time() - start_time > timeout:
                logger.error('[BATCH DEBUG] Analysis timed out')
                raise TimeoutError("Analysis timed out after 30 minutes")
            
            # Update analysis with results
            analysis.accuracy_score = result['statistics']['accuracy_score']
            analysis.total_words = result['statistics']['total_words']
            analysis.correct_words = result['statistics']['correct_words']
            analysis.missing_words = result['statistics']['missing_words']
            analysis.wrong_words = result['statistics']['wrong_words']
            analysis.save()
            logger.info(f'[BATCH DEBUG] Analysis updated with results')
            
            # Create detailed result with better error handling
            try:
                analysis_result, created = AnalysisResult.objects.get_or_create(
                    analysis=analysis,
                    defaults={
                        'transcribed_text': result['transcribed_text'],
                        'script_text': result['script_text'],
                        'processing_time': result['processing_time'],
                        'whisper_model_used': analyzer.model_size,
                        'segments': result.get('segments', None),
                    }
                )

                if not created:
                    analysis_result.transcribed_text = result['transcribed_text']
                    analysis_result.script_text = result['script_text']
                    analysis_result.processing_time = result['processing_time']
                    analysis_result.segments = result.get('segments', None)
                    analysis_result.save()
                logger.info(f'[BATCH DEBUG] AnalysisResult saved for analysis {analysis_id}')
            except Exception as e:
                logger.error(f"[BATCH DEBUG] Failed to save AnalysisResult for analysis {analysis_id}: {e}")
                raise Exception(f"Failed to save analysis results: {e}")
            
            # Create word comparisons
            WordComparison.objects.filter(analysis=analysis).delete()
            for comp in result['comparisons']:
                WordComparison.objects.create(
                    analysis=analysis,
                    script_word=comp['script_word'],
                    audio_word=comp['audio_word'],
                    word_index=comp['word_index'],
                    is_correct=comp['is_correct'],
                    similarity_score=comp['similarity_score'],
                    error_type=comp['error_type']
                )
            logger.info(f'[BATCH DEBUG] Word comparisons created for analysis {analysis_id}')
            
            # Verify AnalysisResult exists before cleanup
            try:
                verification_result = AnalysisResult.objects.get(analysis=analysis)
                if not verification_result.transcribed_text or not verification_result.script_text:
                    logger.error(f"[BATCH DEBUG] AnalysisResult for analysis {analysis_id} is incomplete - skipping cleanup")
                    raise Exception("AnalysisResult is incomplete")
                logger.info(f"[BATCH DEBUG] AnalysisResult verification passed for analysis {analysis_id}")
            except AnalysisResult.DoesNotExist:
                logger.error(f"[BATCH DEBUG] AnalysisResult not found for analysis {analysis_id} - skipping cleanup")
                raise Exception("AnalysisResult not found")
            
            # Only clean up files after AnalysisResult is successfully saved and verified
            cleanup_analysis_files(analysis)
            logger.info(f'[BATCH DEBUG] Cleanup complete for analysis {analysis_id}')
        except Exception as e:
            logger.error(f"[BATCH DEBUG] Error in analysis {analysis_id}: {e}")
            # Mark analysis as failed
            analysis.accuracy_score = -1  # Use -1 to indicate failure
            analysis.save()
            raise
        
    except AudioAnalysis.DoesNotExist:
        logger.error(f"[BATCH DEBUG] Analysis {analysis_id} not found")
    except Exception as e:
        logger.error(f"[BATCH DEBUG] Critical error in analysis {analysis_id}: {e}")

def run_analysis(analysis_id):
    """Legacy function - now calls the timeout version"""
    run_analysis_with_timeout(analysis_id)

def analysis_detail(request, analysis_id):
    print("[DEBUG] analysis_detail called")
    analysis = get_object_or_404(AudioAnalysis, id=analysis_id)
    result = None
    segments = []
    try:
        result = analysis.detailed_result
        print(f"[DEBUG] AnalysisResult found: {result}")
        # Try to load segments from a JSONField if you add it, or fallback to re-transcribe if needed
        if hasattr(result, 'segments') and result.segments:
            segments = result.segments
        else:
            # Optionally, you could re-run the analyzer here to get segments if not stored
            pass
    except Exception as e:
        print(f"[DEBUG] No AnalysisResult: {e}")
    mode = request.GET.get('mode', 'word')
    paragraph_results = None
    if mode == 'paragraph':
        analyzer = AudioAnalyzer()
        if result:
            print(f"[DEBUG] Using AnalysisResult for paragraph analysis")
            script_text = result.script_text
            audio_text = result.transcribed_text
        else:
            print(f"[DEBUG] Reconstructing script/audio text from WordComparison for paragraph analysis")
            comparisons = analysis.word_comparisons.all()
            script_text = ' '.join([c.script_word for c in comparisons if c.script_word])
            audio_text = ' '.join([c.audio_word for c in comparisons if c.audio_word])
        paragraph_results = analyzer.compare_paragraphs(script_text, audio_text, script_path=None, force_single_paragraph=True, segments=segments)
    if analysis.accuracy_score == -1:
        messages.error(request, 'Analysis failed. Please try again with a smaller file or different format.')
        return redirect('audio_checker:home')
    comparisons = analysis.word_comparisons.all()
    paginator = Paginator(comparisons, 50)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    context = {
        'analysis': analysis,
        'page_obj': page_obj,
        'comparisons': page_obj,
        'mode': mode,
        'paragraph_results': paragraph_results,
        'segments': segments,
    }
    return render(request, 'audio_checker/analysis_detail.html', context)

def analysis_list(request):
    """List all analyses"""
    analyses = AudioAnalysis.objects.all().order_by('-created_at')
    paginator = Paginator(analyses, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'audio_checker/analysis_list.html', {'page_obj': page_obj})

@csrf_exempt
def check_analysis_status(request, analysis_id):
    """AJAX endpoint to check analysis status"""
    try:
        analysis = AudioAnalysis.objects.get(id=analysis_id)
        
        # Check if analysis failed
        if analysis.accuracy_score == -1:
            return JsonResponse({
                'status': 'failed',
                'error': 'Analysis failed. Please try again.'
            })
        
        return JsonResponse({
            'status': 'completed' if analysis.accuracy_score is not None else 'processing',
            'accuracy': analysis.accuracy_score,
            'total_words': analysis.total_words,
            'correct_words': analysis.correct_words,
            'missing_words': analysis.missing_words,
            'wrong_words': analysis.wrong_words,
        })
    except AudioAnalysis.DoesNotExist:
        return JsonResponse({'error': 'Analysis not found'}, status=404)

def download_results(request, analysis_id):
    """Download analysis results as CSV"""
    analysis = get_object_or_404(AudioAnalysis, id=analysis_id)
    
    import csv
    from django.http import HttpResponse
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="analysis_{analysis_id}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Word Index', 'Script Word', 'Audio Word', 'Error Type', 'Similarity Score', 'Is Correct'])
    
    for comparison in analysis.word_comparisons.all():
        writer.writerow([
            comparison.word_index,
            comparison.script_word,
            comparison.audio_word,
            comparison.error_type,
            comparison.similarity_score,
            comparison.is_correct
        ])
    
    return response

def download_transcript_docx(request, analysis_id):
    """Download the transcribed audio as a DOCX file"""
    analysis = get_object_or_404(AudioAnalysis, id=analysis_id)
    
    # Check if analysis has results
    if analysis.accuracy_score == -1:
        messages.error(request, 'Analysis failed. No transcript available.')
        return redirect('audio_checker:analysis_detail', analysis_id=analysis_id)
    
    try:
        # Get the transcribed text
        transcribed_text = ""
        try:
            # Try to get from AnalysisResult first
            result = analysis.detailed_result
            transcribed_text = result.transcribed_text
        except:
            # Fallback: reconstruct from word comparisons
            comparisons = analysis.word_comparisons.all()
            transcribed_text = ' '.join([c.audio_word for c in comparisons if c.audio_word])
        
        if not transcribed_text.strip():
            messages.error(request, 'No transcript available for this analysis.')
            return redirect('audio_checker:analysis_detail', analysis_id=analysis_id)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Transcribed Audio: {analysis.title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add analysis info
        doc.add_paragraph(f'Analysis ID: {analysis.id}')
        doc.add_paragraph(f'Created: {analysis.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Accuracy Score: {analysis.accuracy_score:.2f}%')
        doc.add_paragraph(f'Total Words: {analysis.total_words}')
        doc.add_paragraph(f'Correct Words: {analysis.correct_words}')
        doc.add_paragraph(f'Missing Words: {analysis.missing_words}')
        doc.add_paragraph(f'Wrong Words: {analysis.wrong_words}')
        
        # Add a separator
        doc.add_paragraph('')
        doc.add_paragraph('=' * 50)
        doc.add_paragraph('')
        
        # Add the transcribed text
        doc.add_heading('Transcribed Text', level=1)
        
        # Split text into paragraphs and add them
        paragraphs = transcribed_text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Create the response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="transcript_{analysis_id}_{analysis.title.replace(" ", "_")}.docx"'
        
        # Save the document to the response
        doc.save(response)
        
        return response
        
    except Exception as e:
        logger.error(f"Error creating DOCX transcript for analysis {analysis_id}: {e}")
        messages.error(request, 'Error creating transcript file. Please try again.')
        return redirect('audio_checker:analysis_detail', analysis_id=analysis_id)

def transcript_segments_view(request, analysis_id):
    import json
    from .services import AudioAnalyzer
    analysis = get_object_or_404(AudioAnalysis, id=analysis_id)
    result = None
    segments = []
    warning = None
    
    # First try to get segments from AnalysisResult
    try:
        result = analysis.detailed_result
        if hasattr(result, 'segments') and result.segments:
            segments = result.segments
            if isinstance(segments, str):
                segments = json.loads(segments)
    except Exception as e:
        pass
    
    # Fallback: if segments are missing, try to re-run transcription
    if not segments:
        try:
            # Check if audio file exists
            if not analysis.audio_file or not os.path.exists(analysis.audio_file.path):
                warning = f"Audio file not found: {analysis.audio_file.name if analysis.audio_file else 'No file'}. Cannot regenerate segments."
            else:
                analyzer = AudioAnalyzer()
                audio_path = analysis.audio_file.path
                _, _, segments = analyzer.transcribe_audio(audio_path)
                warning = "Segments were missing and have been regenerated. This may take extra time."
        except Exception as e:
            warning = f"Could not regenerate segments: {e}"
    
    print("Segments in view:", segments)
    context = {
        'analysis': analysis,
        'segments': segments,
        'warning': warning,
    }
    return render(request, 'audio_checker/transcript_segments.html', context)

def cleanup_analysis_files(analysis):
    """Clean up uploaded files after analysis is complete"""
    try:
        if analysis.script_file and os.path.exists(analysis.script_file.path):
            os.remove(analysis.script_file.path)
        if analysis.audio_file and os.path.exists(analysis.audio_file.path):
            os.remove(analysis.audio_file.path)
        logger.info(f"Cleaned up files for analysis {analysis.id}")
    except Exception as e:
        logger.error(f"Error cleaning up files for analysis {analysis.id}: {e}")

# New batch upload views
def batch_upload(request):
    """Handle batch upload of multiple script-audio pairs"""
    import logging
    logger = logging.getLogger(__name__)
    if request.method == 'POST':
        print("=== BATCH UPLOAD DEBUG ===")
        print(f"POST data keys: {list(request.POST.keys())}")
        print(f"FILES data keys: {list(request.FILES.keys())}")
        logger.info('Batch upload POST received')
        batch_form = BatchUploadForm(request.POST)
        if batch_form.is_valid():
            print("Batch form is valid")
            logger.info('Batch form is valid')
            batch = batch_form.save(commit=False)
            if request.user.is_authenticated:
                batch.user = request.user
            batch.save()
            print(f"Batch created: {batch}")
            logger.info(f'Batch created: {batch}')
            
            # Process the uploaded files
            files_data = request.POST.get('files_data')
            print(f"files_data raw: {files_data}")
            logger.info(f'files_data raw: {files_data}')
            if files_data:
                try:
                    files_list = json.loads(files_data)
                    print(f"files_list parsed: {files_list}")
                    logger.info(f'files_list parsed: {files_list}')
                    for file_data in files_list:
                        if file_data.strip():
                            try:
                                data = json.loads(file_data)
                                title = data.get('title', 'Untitled Analysis')
                                pair_id = data.get('pair_id', '1')
                                logger.info(f'Processing pair_id={pair_id}, title={title}')
                                
                                # Get the actual file objects from request.FILES
                                script_file_key = f"script_file_{pair_id}"
                                audio_file_key = f"audio_file_{pair_id}"
                                logger.info(f'Looking for files: {script_file_key}, {audio_file_key}')
                                script_file = request.FILES.get(script_file_key)
                                audio_file = request.FILES.get(audio_file_key)
                                logger.info(f'Got script_file: {script_file}, audio_file: {audio_file}')
                                
                                if script_file and audio_file:
                                    print(f"Creating AudioAnalysis for pair_id={pair_id}")
                                    analysis = AudioAnalysis.objects.create(
                                        title=title,
                                        script_file=script_file,
                                        audio_file=audio_file,
                                        user=request.user if request.user.is_authenticated else None,
                                        batch=batch
                                    )
                                    print(f"Created AudioAnalysis: {analysis}")
                                    logger.info(f'Created AudioAnalysis: {analysis}')
                                    
                                    # Start analysis in background
                                    thread = threading.Thread(target=run_analysis_with_timeout, args=(analysis.id,))
                                    thread.daemon = True
                                    thread.start()
                                    print(f"Started analysis thread for {analysis.id}")
                                    logger.info(f'Started analysis thread for {analysis.id}')
                                else:
                                    print(f"Missing files for pair_id={pair_id}")
                                    logger.error(f'Missing files for pair_id={pair_id}')
                            except Exception as e:
                                logger.error(f"Error processing file pair: {e}")
                                continue
                except json.JSONDecodeError:
                    logger.error("Invalid JSON in files_data")
            else:
                logger.error('No files_data found in POST')
            
            messages.success(request, f'Batch upload "{batch.title}" created with {batch.get_total_analyses()} analyses!')
            return redirect('audio_checker:batch_detail', batch_id=batch.id)
        else:
            logger.error('Batch form is invalid')
    else:
        batch_form = BatchUploadForm()
    
    return render(request, 'audio_checker/batch_upload.html', {'batch_form': batch_form})

def batch_detail(request, batch_id):
    """Show details of a batch upload"""
    batch = get_object_or_404(BatchUpload, id=batch_id)
    analyses = batch.analyses.all().order_by('-created_at')
    
    # Update batch status based on analyses
    total = batch.get_total_analyses()
    completed = batch.get_completed_analyses()
    failed = batch.get_failed_analyses()
    
    if total == 0:
        batch.status = 'pending'
    elif completed + failed == total:
        batch.status = 'completed' if failed == 0 else 'failed'
    else:
        batch.status = 'processing'
    batch.save()
    
    context = {
        'batch': batch,
        'analyses': analyses,
        'total': total,
        'completed': completed,
        'failed': failed,
        'pending': total - completed - failed
    }
    return render(request, 'audio_checker/batch_detail.html', context)

def batch_list(request):
    """List all batch uploads"""
    batches = BatchUpload.objects.all().order_by('-created_at')
    paginator = Paginator(batches, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'audio_checker/batch_list.html', {'page_obj': page_obj})

@csrf_exempt
def upload_file_pair(request):
    """AJAX endpoint to handle individual file pair uploads"""
    if request.method == 'POST':
        try:
            title = request.POST.get('title')
            script_file = request.FILES.get('script_file')
            audio_file = request.FILES.get('audio_file')
            
            if not all([title, script_file, audio_file]):
                return JsonResponse({'error': 'All fields are required'}, status=400)
            
            # Validate file types
            if not script_file.name.endswith(('.docx', '.doc')):
                return JsonResponse({'error': 'Invalid script file format'}, status=400)
            
            if not audio_file.name.endswith(('.wav', '.mp3', '.m4a', '.flac')):
                return JsonResponse({'error': 'Invalid audio file format'}, status=400)
            
            return JsonResponse({
                'success': True,
                'title': title,
                'script_file': script_file.name,
                'audio_file': audio_file.name
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)
